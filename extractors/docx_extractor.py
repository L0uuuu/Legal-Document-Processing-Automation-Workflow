"""Extract text from Microsoft Word (.docx) documents."""

import time

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from models.extraction import (
    DocumentFormat,
    ExtractionMethod,
    PageContent,
    ExtractionResult,
)
from extractors.base import BaseExtractor


class DOCXExtractor(BaseExtractor):
    """
    Extracts text from .docx files using python-docx.

    Note: DOCX files don't have a native "page" concept,
    so we treat the entire document as a single page,
    or split by sections/page breaks if present.
    """

    def supported_formats(self) -> list[DocumentFormat]:
        return [DocumentFormat.DOCX]

    def extract(self, file_path: str) -> ExtractionResult:
        self.validate_file(file_path)
        start_time = time.time()

        pages = self._extract_docx(file_path)

        return self._build_result(
            file_path=file_path,
            file_format=DocumentFormat.DOCX,
            extraction_method=ExtractionMethod.DOCX_PARSE,
            pages=pages,
            start_time=start_time,
        )

    def _extract_docx(self, file_path: str) -> list[PageContent]:
        """Extract text from DOCX, splitting by section breaks."""
        try:
            doc = Document(file_path)
        except PackageNotFoundError:
            self.warnings.append("File is not a valid .docx package.")
            return []
        except Exception as e:
            self.warnings.append(f"Error opening DOCX: {str(e)}")
            return []

        sections: list[list[str]] = [[]]
        current_section = 0

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()

            # Detect page/section breaks
            if self._has_page_break(paragraph):
                sections.append([])
                current_section += 1

            if text:
                sections[current_section].append(text)

        # Also extract text from tables
        table_texts = self._extract_tables(doc)
        if table_texts:
            # Append table content to the last section
            sections[-1].extend(table_texts)

        # Build page objects
        pages = []
        for idx, section_paragraphs in enumerate(sections):
            text = "\n".join(section_paragraphs)
            if not text.strip():
                continue

            languages = self.detect_languages(text)
            pages.append(
                PageContent(
                    page_number=idx + 1,
                    text=self.clean_text(text),
                    has_arabic=languages["has_arabic"],
                    has_french=languages["has_french"],
                )
            )

        return pages

    def _has_page_break(self, paragraph) -> bool:
        """Check if a paragraph contains a page break."""
        for run in paragraph.runs:
            # Check for explicit page break in run XML
            if run._element.xml and "w:br" in run._element.xml:
                if 'w:type="page"' in run._element.xml:
                    return True
        return False

    def _extract_tables(self, doc: Document) -> list[str]:
        """Extract text content from all tables in the document."""
        table_texts = []

        for table in doc.tables:
            rows_text = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    rows_text.append(" | ".join(cells))

            if rows_text:
                table_texts.append("\n".join(rows_text))

        return table_texts